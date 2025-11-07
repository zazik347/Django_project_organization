from datetime import datetime
from io import BytesIO

from django.db.models.functions import Cast, Coalesce
from django.utils import timezone
from docx import Document
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from avtomatizations.models import AutomationEngineer
from projects.models import Project, ProjectEmployee, ProjectEquipment
from django.db.models import ExpressionWrapper, Sum, fields, F, DecimalField, DurationField, IntegerField, Case, When, \
    Value


def project_employees(request, pk):
    project = get_object_or_404(Project, pk=pk)
    # Получаем всех сотрудников с ролями из модели связи
    employees_with_roles = ProjectEmployee.objects.filter(project=project)
    project_equipment = ProjectEquipment.objects.filter(project=project)
    equipment = project.equipment.all()
    # Получаем всех подрядчиков через связь
    subcontractors = project.subcontractors.all()
    return render(request, 'project_employees.html', {
        'project': project,
        'employees_with_roles': employees_with_roles,
        'equipment': equipment,
        'subcontractors': subcontractors,
    })

def project_equipment(request, pk):
    project = get_object_or_404(Project, pk=pk)
    equipment = project.equipment.all()  # Все оборудование этого проекта
    return render(request, 'project_equipment.html', {
        'project': project,
        'equipment': equipment
    })
def project_list(request):
    projects = Project.objects.all()
    return render(request, 'project_list.html', {
        'projects': projects
    })

def export_project_to_word(request, project_id):
    # Получаем проект по ID
    global table, hdr_cells
    project = get_object_or_404(Project, id=project_id)
    employees_with_roles = ProjectEmployee.objects.filter(project=project)
    equipment = project.equipment.all()
    subcontractors = project.subcontractor_assignments.all()
    print(subcontractors)

    # Создаём новый документ Word
    doc = Document()

    # Добавляем заголовок
    doc.add_heading(f'Информация о проекте: {project.title}', level=1)

    # Добавляем Менаджера
    if project.manager.full_name:
        doc.add_heading(f'Менеджер: {project.manager.full_name}', level=1)
    else:
        doc.add_paragraph('Менаджер  отсутствует.')

    # Добавляем дату начала
    if project.start_date:
        doc.add_heading(f'Дата начала: {project.start_date}', level=1)
    else:
        doc.add_paragraph('Дата начала отсутствует.')

    # Добавляем дату окончания
    if project.end_date:
        doc.add_heading(f'Дата окончания: {project.end_date}', level=1)
    else:
        doc.add_paragraph('Дата окончания отсутствует.')

    # Добавляем стоимость проекта
    if project.cost:
        doc.add_heading(f'Стоимость: {project.cost}', level=1)
    else:
        doc.add_paragraph('Стоимость.')

    # Добавляем описание проекта
    if project.description:
        doc.add_heading(f'Описание проекта:')
        doc.add_paragraph('')
        doc.add_paragraph({project.description})
    else:
        doc.add_paragraph('Описание отсутствует.')

        # Добавляем раздел "Сотрудники"
    doc.add_heading('Сотрудники:', level=2)
    doc.add_paragraph('')

    # Если есть сотрудники, добавляем их в таблицу
    if employees_with_roles .exists():
        table = doc.add_table(rows=1, cols=3)
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Проверяем, есть ли runs
            if not paragraph.runs:
                # Добавляем пустой run
                paragraph.add_run()

            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(12)

        hdr_cells[0].text = 'ФИО'
        hdr_cells[1].text = 'Отдел'
        hdr_cells[2].text = 'Должность'

        # Выравниваем текст по центру
        for cell in hdr_cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавляем строки сотрудников
        for assignment in employees_with_roles.all():
            row_cells = table.add_row().cells

            # ФИО
            row_cells[0].text = assignment.employee.full_name
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Отдел
            if hasattr(assignment.employee, 'department'):
                row_cells[1].text = assignment.employee.department.name
            else:
                row_cells[1].text = ''
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Должность
            row_cells[2].text = assignment.role
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Устанавливаем границы у ячеек
            for cell in row_cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border in ['top', 'left', 'bottom', 'right']:
                    bdr = OxmlElement(f'w:{border}')
                    bdr.set(qn('w:val'), 'single')
                    bdr.set(qn('w:sz'), '4')  # толщина границы
                    bdr.set(qn('w:color'), '888888')  # цвет границы
                    tcBorders.append(bdr)
                tcPr.append(tcBorders)
    else:
        doc.add_paragraph('Нет назначенных сотрудников.')

    # Добавляем раздел "Оборудование"
    doc.add_heading('Оборудование:', level=2)
    doc.add_paragraph('')
    # Если есть оборудование, добавляем их в таблицу
    if equipment .exists():
        table = doc.add_table(rows=1, cols=1)
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Проверяем, есть ли runs
            if not paragraph.runs:
                # Добавляем пустой run
                paragraph.add_run()

            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(12)

        hdr_cells[0].text = 'Название оборудования'

        # Выравниваем текст по центру
        for cell in hdr_cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Добавляем список оборудования
        for item in equipment.all():
            row_cells = table.add_row().cells
            # Оборудование
            row_cells[0].text = item.equipment.name
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Устанавливаем границы у ячеек
            for cell in row_cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border in ['top', 'left', 'bottom', 'right']:
                    bdr = OxmlElement(f'w:{border}')
                    bdr.set(qn('w:val'), 'single')
                    bdr.set(qn('w:sz'), '4')  # толщина границы
                    bdr.set(qn('w:color'), '888888')  # цвет границы
                    tcBorders.append(bdr)
                tcPr.append(tcBorders)
    else:
        doc.add_paragraph('Оборудование не испльзавалось.')
    # Добавляем раздел "Подрядчик"
    doc.add_heading('Подрядчики:', level=2)
    doc.add_paragraph('')
    # Если есть Подрядчики, добавляем их в таблицу
    if subcontractors.exists():
        table = doc.add_table(rows=1, cols=2)
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Проверяем, есть ли runs
            if not paragraph.runs:
                # Добавляем пустой run
                paragraph.add_run()

            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(12)

        hdr_cells[0].text = 'Название подрядчика'
        hdr_cells[1].text = 'Стоимость услуг'


        # Выравниваем текст по центру
        for i, cell in enumerate(hdr_cells):
            paragraph = cell.paragraphs[0]
            if i == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if i == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Добавляем список подрядчиков
        for sub in subcontractors:
            row_cells = table.add_row().cells
            # Название подрядчика
            row_cells[0].text = sub.subcontractor.name
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Стоимость услуг
            row_cells[1].text = str(sub.cost)
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Устанавливаем границы у ячеек
            for cell in row_cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border in ['top', 'left', 'bottom', 'right']:
                    bdr = OxmlElement(f'w:{border}')
                    bdr.set(qn('w:val'), 'single')
                    bdr.set(qn('w:sz'), '4')  # толщина границы
                    bdr.set(qn('w:color'), '888888')  # цвет границы
                    tcBorders.append(bdr)
                tcPr.append(tcBorders)
    else:
        doc.add_paragraph('Подрядчики не задействованы.')
    # Сохраняем документ в байтах
    f = BytesIO()
    doc.save(f)
    length = f.tell()
    f.seek(0)

    # Возвращаем ответ как скачиваемый файл
    response = HttpResponse(f, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="project_{project.id}.docx"'

    return response

def active_projects(request):
    # Текущие проекты — те, которые сейчас активны (end_date ещё не наступила)
    current_date = timezone.now().date()
    active_projects = Project.objects.filter(start_date__lte=current_date, end_date__gte=current_date)

    return render(request, 'active_projects.html', {
        'active_projects': active_projects,
    })
def projects_by_date(request):
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')

    # Если дата указана — конвертируем её
    try:
        if start_date:
            start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
        if end_date:
            end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
    except ValueError:
        start_date = None
        end_date = None

    # Фильтруем проекты
    if start_date and end_date:
        filtered_projects = Project.objects.filter(
            start_date__gte=start_date,
            end_date__lte=end_date
        )
    elif start_date:
        filtered_projects = Project.objects.filter(start_date__gte=start_date)
    elif end_date:
        filtered_projects = Project.objects.filter(end_date__lte=end_date)
    else:
        filtered_projects = Project.objects.all()

    return render(request, 'projects_by_date.html', {
        'filtered_projects': filtered_projects,
        'start_date': start_date,
        'end_date': end_date,
    })

def search_project_detail(request, project_id):
    # Получаем проект по ID
    project = get_object_or_404(Project, id=project_id)
    # Получаем договор, к которому относится этот проект
    contract = project.contract

    return render(request, 'search_project_detail.html', {
        'project': project,
        'contract': contract,
    })

def project_efficiency(request):
    # Шаг 1: Вычисляем продолжительность проекта в днях
    duration_in_seconds = ExpressionWrapper(
        F('end_date') - F('start_date'),
        output_field=DurationField()
    )
    duration_in_days = ExpressionWrapper(
        Cast(duration_in_seconds / 86400, output_field=IntegerField()),
        output_field=IntegerField()
    )

    # Аннотируем total_cost, заменяя NULL на 0
    projects = Project.objects.annotate(
        total_cost=Coalesce(
            ExpressionWrapper(
                Sum(
                    F('employee_assignments__employee__hourly_rate') * 8 * duration_in_days
                ),
                output_field=DecimalField(max_digits=15, decimal_places=2)
            ),
            Value(0.0),
            output_field=DecimalField(max_digits=15, decimal_places=2)
        )
    )

    # Шаг 3: Аннотируем efficiency с проверкой на zero
    projects = projects.annotate(
        efficiency=Case(
            When(total_cost__gt=0, then=ExpressionWrapper(
                ((F('cost') - F('total_cost')) / F('total_cost')) * 100,
                output_field=DecimalField(max_digits=15, decimal_places=2)
            )),
            default=Value(100.00),  # если total_cost == 0 → эффективность 100%
            output_field=DecimalField(max_digits=15, decimal_places=2)
        )
    ).order_by('-total_cost')


    for p in projects:
        print(f"Проект: {p.title}, Затраты: {p.total_cost}, Эффективность: {p.efficiency}")

    return render(request, 'project_efficiency.html', {
        'projects': projects
    })
